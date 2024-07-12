import fitz  # PyMuPDF
import os
import sys
from pdf2image import convert_from_path
from google.cloud import vision
import io
import re
from docx import Document
import pytesseract
import json
from PIL import Image

os.environ['HF_HOME'] = "cache"
poppler_path = r'poppler-24.02.0/Library/bin'

pytesseract.pytesseract.tesseract_cmd = "Tesseract-OCR/tesseract.exe"

from transformers import pipeline

"""****************************************JSON FILE EMBEDDED********************************************"""
# Set the environment variable for Google Application Credentials

import tempfile
from google.cloud import vision_v1
service_account_info = {
  "type": "",
  "project_id": "",
  "private_key_id": "",
  "private_key": "",
  "client_email": "",
  "client_id": "",
  "auth_uri": "",
  "token_uri": "",
  "auth_provider_x509_cert_url": "",
  "client_x509_cert_url": ""
  "universe_domain": ""
}
def create_temp_credentials_file():
    # Create a temporary file and write the JSON data to it
    with tempfile.NamedTemporaryFile(delete=False, mode='w') as temp_file:
        json.dump(service_account_info, temp_file)
        temp_file_path = temp_file.name
    return temp_file_path

temp_credentials_path = create_temp_credentials_file()
os.environ['GOOGLE_APPLICATION_CREDENTIALS'] = temp_credentials_path
#os.environ['GOOGLE_APPLICATION_CREDENTIALS'] = r'qualified-star-428805-i7-d1f63f348f51.json'    

"""***************************************************************************************************************"""

# Set up Google Cloud Vision client
client = vision.ImageAnnotatorClient()

#detecting text from images
def detect_text(image_path):
    """Detects text in the file."""
    with io.open(image_path, 'rb') as image_file:
        content = image_file.read()

    image = vision.Image(content=content)

    response = client.text_detection(image=image)
    texts = response.text_annotations

    if response.error.message:
        raise Exception(f'{response.error.message}')
    #print(texts)
    return texts[0].description if texts else "No text found"


#iterating over PDF pages
def extract_text_from_pdf(file_path):
    text = ""
    pdf_document = fitz.open(file_path)
    num_pages = len(pdf_document)
    print("extracting text from pdf")
    if num_pages > 15:
        for page_num in range(1, 16, 1):  
            page = pdf_document.load_page(page_num - 1)
            text += page.get_text()
    else:
        for page_num in range(num_pages):
            page = pdf_document.load_page(page_num)
            text += page.get_text()
    
    pdf_document.close()
    #print(text)
    return text


def extract_text_from_images(file_path):
    images = convert_from_path(file_path)
    text = ""
    print("extracting text from image")
    for idx, image in enumerate(images):
        image_path = f"temp_page_{idx}.png"
        image.save(image_path, 'PNG')
        text += detect_text(image_path)
        os.remove(image_path)  # Clean up the temporary image file
    #print(text)
    return text

"""*************** EXTRACTING YEARS ************************"""

def extract_years(ext_txt):
    try:
        date_pattern = r'\b(?:\d{1,2}[/-]\d{1,2}[/-])(\d{2,4})\b'
        dates = re.findall(date_pattern, ext_txt)
        
        if not dates:
            return ['x']  # Return 'x' if no dates are found

        years = [date[-2:] for date in dates if len(date) >= 2]  # Extract the last two digits of the year if the date has enough characters
        
        if not years:
            return ['x']  # Return 'x' if the years list is empty

        return years
    except Exception as e:
        print(f"Error extracting years: {e}")
        return ['x']
    
"""***********************************************************"""
carrier_names = [
    "AIG",
    "All Risks Limited",
    "Allstate",
    "American Modern Insurance",
    "AmGUARD",
    "AmTrust Financial Services, Inc.",
    "AmWINS Access Insurance",
    "AmWINS Brokerage",
    "Apogee Insurance Group",
    "Arch Insurance Company",
    "Arthur J. Gallagher",
    "ASI",
    "Attune",
    "Bass Underwriters",
    "Berkley Asset Protection",
    "Berkley Regional Specialty Insurance Company",
    "Berkshire Hathaway Inc.",
    "Bold Penguin",
    "Breckenridge Insurance Services",
    "Burns & Wilcox",
    "Chesapeake Employers Insurance",
    "CannGen Insurance Services",
    "Chubb",
    "CID Insurance Programs",
    "Cincinnati Insurance",
    "Clearcover Insurance Company",
    "Clearcover, Inc.",
    "CNA Surety",
    "Colonial Surety Company",
    "COVER Financial",
    "Cowbell Insurance",
    "Crum & Forster Specialty Insurance Company",
    "CRC Group Wholesale & Specialty",
    "Donegal",
    "ECC Insurance Brokers, LLC",
    "Employers Insurance Company of Nevada",
    "Encompass Insurance",
    "Everguard Insurance",
    "FIRST CONNECT",
    "FISIF-Food Industry Self-Insurance Fund",
    "Foremost",
    "Front Row Insurance",
    "GIG Insurance Group",
    "Hartford",
    "HBW"
    "Hiscox",
    "Hull & Company",
    "IBIS-Independent Bankers Insurance Services",
    "IIAA Agency Administrative Services, Inc.",
    "Insential Inc.",
    "Insurance Solutions Group",
    "International Excess",
    "Irwin Siegel Agency, Inc.",
    "Kinsale Insurance",
    "Kelly Klee",
    "Kemper",
    "Kentucky Employers' Mutual Insurance",
    "Liberty Mutual",
    "Mesa Underwriters Specialty Insurance",
    "Markel Insurance Company",
    "Mercer Insurance Group",
    "Mercury County Mutual - TX",
    "Mercury Insurance",
    "MetLife Auto & Home®",
    "Mountain States Mutual Casualty Company",
    "National Fire Insurance of Hartford",
    "National Program Administrators",
    "Nationwide Insurance",
    "Neptune",
    "NETWORKED | AMWINS DAIS",
    "New Heights Insurance Solutions",
    "NEW MEXICO FAIR PLAN",
    "New Mexico Fair Plan",
    "New Mexico Mutual Casualty Company",
    "NEXT Insurance",
    "NFP Property & Casualty Services, Inc.",
    "Ohio Casualty",
    "One Beacon",
    "Orchid",
    "Pathpoint",
    "Penn National Insurance",
    "PEACHTREE SPECIAL RISK BROKERS, LLC",
    "Philadelphia Insurance Company",
    "PIE Insurance",
    "Preferred Employers Insurance Company",
    "Privilege Underwriters Reciprocal Exchange",
    "Progressive Insurance",
    "PROPELLER BONDS",
    "QBE Insurance Corporation",
    "Redpoint County Mutual Company",
    "Risk Exchange Insurance Services, Inc.",
    "Risk Placement Service",
    "RLI Insurance Company",
    "RLI Surety",
    "Ryan Turner Specialty",
    "Starstone National Insurance",
    "Safeco Insurance",
    "SELECTIVE INSURANCE GROUP INC.",
    "Smart Choice Agents",
    "Southwest Insurance Group",
    "SPECIAL MARKETS INSURANCE CONSULTANTS",
    "Specialty Market Insurance Consultants, Inc.",
    "State Auto Insurance",
    "Stillwater Insurance Services, Inc.",
    "Surety Placement Services, LLC",
    "TAPCO Underwriters Inc.",
    "Texas Mutual",
    "The Hartford",
    "Thimble Insurance Services",
    "Topa Insurance Company",
    "TransGuard Insurance Company of America, Inc.",
    "Travelers",
    "Travelers Texas MGA",
    "Union Standard Insurance Group",
    "US Assure",
    "USLI",
    "Virtus Insurance",
    "Western Pacific Insurance Network, Inc.",
    "Western Pacific Mutual Insurance Company",
    "Wright National Flood Insurance Company",
    "Zurich"
] 

"""************* Finding carrier name ***********************"""
def find_carrier_name(ext_txt):
    for carrier in carrier_names:
        if re.search(re.escape(carrier), ext_txt, re.IGNORECASE):
            return carrier
    return "UnknownCarrier"
  
"""***********************************************************"""
"""  
def determine_line_of_business(text): 
    lob_keywords = {
    "Auto ": ["Collision", "Comprehensive"],
    "Cyber Liability ": ["Cyber","Data breach", "Cyberattack", "Privacy liability", "Network security"],
    "Directors & Officers ": ["directors","Fiduciary duty", "Shareholder litigation", "Corporate governance", "Indemnification", "Coverage triggers"],
    "Employment Practices Liability ": ["e mployement practices liability","Discrimination", "Harassment", "Retaliation", "Wrongful termination", "Coverage territory"],
    "Errors & Omissions ": ["Professional negligence", "Legal liability", "Malpractice", "Claims-made policy"],
    "Excess Policy ": ["excess policy","Surplus", "Catastrophe", "Reinsurance", "Policy aggregate limit"],
    "General Liability ": ["Bodily injury", "Property damage", "Lawsuit", "Occurrence", "Coverage territory"],
    "Homeowners ": ["homeowners","Dwelling coverage", "Personal property", "Additional living expenses", "Peril coverage"],
    "Individual Health ": ["Health insurance", "Medical coverage", "Pre-existing conditions"],
    "Inland Marine ": ["Transportation", "Cargo", "Logistics", "Bailee coverage"],
    "Professional Liability ": ["Errors", "Omissions", "Negligence", "Professional services", "Claims-made trigger"],
    "Property ": ["property","Real estate", "Buildings", "Contents", "Peril coverage", "Replacement cost"],
    "Umbrella ": ["umbrella","Catastrophic", "Excess", "Liability", "Aggregate limit", "Underlying coverage"],
    "Flood ": ["flood","Water damage", "Natural disaster", "Floodplain", "Elevation certificate", "Coverage waiting period"],
    "Dwelling Fire ": ["Residential property", "Fire damage", "Rental dwelling", "Named perils", "Vacancy clause"],
    "Mobile Home ": ["Manufactured housing", "Trailer", "Park model", "Personal property coverage", "Mobile home park endorsement"],
    "Businessowners Policy ": ["Business property", "Business interruption", "Liability", "Business personal property", "Building ordinance coverage"],
    "Boat ": ["Watercraft", "Marine insurance", "Hull insurance", "Navigational territory", "Agreed value coverage"],
    "Earthquake ": ["Earthquake coverage", "Aftershocks", "Seismic retrofitting"],
    "Garage ": ["Repair shop", "Auto service", "Mechanic liability", "Garage liability", "Customer vehicles"],
    "Fire ": ["Blaze", "Conflagration", "Arson", "Fire department service clause", "Fire insurance maps"],
    "Bond ": ["Surety", "Bail", "Performance bond", "Bond premium"],
    "Workers Compensation ": ["Workplace injury", "Disability", "Compensation", "Employer liability", "Medical benefits"]
    }
    keywords_list = [keyword for keywords in lob_keywords.values() for keyword in keywords]
    #print(keywords_list)

    for keyword in keywords_list:
        if re.search(r'\b' + re.escape(keyword) + r'\b', text, re.IGNORECASE):
            for lob, keywords in lob_keywords.items():
                if keyword in keywords:
                    return lob, keyword
    return None, "No line of business found"

"""
def get_policy_type(pdf_path):
    images = convert_from_path(pdf_path, poppler_path=poppler_path, first_page=0, last_page=15)
    for i, image in enumerate(images):
       image.save(f'images/page_{i}.png', 'PNG')
  
    noOfPages = i+1
    
    data = pipeline("document-question-answering",model="impira/layoutlm-document-qa")
    #print("no of images", i)
    base_path = "images/page_{i}.png"
    #questions array
    questions = ["which policy?"]
    results = {}
    
    for numbering in range(noOfPages):
       dynamic_path = base_path.format(i = numbering)
    
       page_results = []

       for ques in questions:
           ans = data(dynamic_path, ques)
           page_results.append({"question" : ques, "answer" : ans})
       results[numbering] = page_results
       print(page_results)
       print("page : ", numbering)

    output_file = "fetched_data.json"
    with open(output_file, "w") as f:
        json.dump(results, f, indent=4)


    input_file = "fetched_data.json"
    with open(input_file, "r") as f:
         data = json.load(f)

    extracted_policy_type = []
    for i, items in data.items():
        for item in items:
            question = item["question"]

            if "answer" in item and item["answer"]:
                extracted_policy_type.append(item["answer"][0]["answer"])
    print(extracted_policy_type)

    policy_type = ["Auto","Cyber Liability","Directors & Officers","Employment Practices Liability","Errors & Omissions","Excess Policy","General Liability","Homeowners","Individual Health","Inland Marine","Professional Liability","Property","Umbrella","Flood","Dwelling Fire","Mobile Home","Businessowners Policy","Boat","Earthquake","Garage","Fire","Bond","Workers Compensation"]

    ans = None

    for extracted in extracted_policy_type:
        for policy in policy_type:
            # Check if any part of policy is in extracted or extracted is in policy
            if policy.lower().replace("'", "") in extracted.lower().replace("'", "") or extracted.lower().replace("'", "") in policy.lower().replace("'", ""):
                ans = policy
                break
        if ans:
            break

    # print("current extracted policy type : ",ans)
    return ans

    # highest_scores = {
    #     "question1": {"score": -1, "answer": None},
    #     "question2": {"score": -1, "answer": None}
    # }
    # # Process each page
    # for page_key, page_data in data.items():
    #     # best_score_question1 = -1
    #     # best_answer_question1 = None
    #     # best_score_question2 = -1
    #     # best_answer_question2 = None

    #     for quesAns, item in enumerate(page_data):
    #         question = item["question"]
        
    #         try:
    #             # Check if answers exist and is not empty
    #             if "answer" in item and item["answer"]:
    #                 answer_data = item["answer"][0]
    #                 score = answer_data["score"]
    #                 answer = answer_data["answer"]
                    
    #                 if score < 1:
    #                     if quesAns == 0:  # First question
    #                         if score > highest_scores["question1"]["score"]:
    #                             highest_scores["question1"]["score"] = score
    #                             highest_scores["question1"]["answer"] = answer
    #                     elif quesAns == 1:  # Second question
    #                         if score > highest_scores["question2"]["score"]:
    #                             highest_scores["question2"]["score"] = score
    #                             highest_scores["question2"]["answer"] = answer

    #             else:
    #                  continue
        
    #         except Exception as e:
    #             #print(f"Error processing question '{question}': {e}")
    #             continue

    # # if best_answer_question1 is not None:         
    # #     highest_scores[page_key + "_Q1"] = {"score": best_score_question1,"answer": best_answer_question1}  
    # #     #print(f"{highest_scores[page_key + "_Q1"]}")  
    # # if best_answer_question2 is not None:         
    # #     highest_scores[page_key +"_Q2"] = {"score": best_score_question2,"answer": best_answer_question2}
    # #     #print(f"{highest_scores[page_key + "_Q2"]}") 

    # # print(highest_scores)
    # hs_policy_type = highest_scores["question1"]["answer"]
    # hs_insurance_company = highest_scores["question2"]["answer"]
    # print(hs_policy_type, hs_insurance_company)

    # return hs_insurance_company, hs_policy_type



    # Print answer with the highest score
    # for question, data in highest_scores.items():
    #     print(f"{data["answer"]}")
    #     print("working")

    # newData = data["answer"].replace(" ", "").lower()
    
    # lob = ['All Lines', 'Auto', 'Cyber', 'Dental', 'Directors & Officers', 'Disability', 'Employment Practices Liability', 'Errors & Omissions', 'Excess Policy', 'General Liability', 'Group Health', 'Group Health Individual', 'Health - Ancillary', 'Home', 'Individual Health', 'Inland Marine', 'Life - Group', 'Life - Individual', 'Life - Voluntary', 'Medicare', 'Professional Liability', 'Property', 'Umbrella', 'Vision', 'Workers Compensation']
    # for i in range(len(lob)):
    #     ans = lob[i].replace(" ", "").lower() in newData
    #     if(ans):
    #         return lob[i]

    # """
    # toBeType = ["home", "auto", "property","general"]
    # extractedType = data['answer']

    # for i in range(len(toBeType)):
    #     ans = toBeType[i] in newData
    #     if(ans):
    #        print(toBeType[i]+'_policy')
   
    # """
    # return newData

    
lob_abbreviations = {
    "All Lines": "All",
    "Auto Liability": "Auto",
    "Cyber Liability": "Cyber",
    "Dental": "Den",
    "Directors & Officers": "D&O",
    "Disability": "DIS",
    "Employment Practices Liability": "EPLI",
    "Errors & Omissions": "E&O",
    "Excess Policy": "XS",
    "General Liability": "GL",
    "Group Health": "GH",
    "Group Health Individual": "GHI",
    "Health - Ancillary": "HA",
    "Homeowners": "HO",
    "Individual Health": "IFP",
    "Inland Marine": "IM",
    "Life - Group": "LG",
    "Life - Individual": "LI",
    "Life - Voluntary": "LV",
    "Medicare": "MED",
    "Professional Liability": "PL",
    "Property": "Prop",
    "Umbrella": "UMB",
    "Vision": "VIS",
    "Workers Compensation": "WC"
}

def get_abbreviated_lob(lob):
    return lob_abbreviations.get(lob, "Unknown")

def process_pdfs_in_folder(folder_path):
    document = Document()
    document.add_heading('Old vs New PDF Names', level=1)
    i=1
    for filename in os.listdir(folder_path):
        if filename.endswith(".pdf") or filename.endswith(".PDF") or filename.endswith(".Pdf") :  # Check if the file is a PDF
            file_path = os.path.join(folder_path, filename)
            print(file_path)
            print("processing the folder")
            print(i)
            print(f"Processing {filename}...")

            # Extract text using PyMuPDF
            pdf_text = extract_text_from_pdf(file_path)

            # If the extracted text is empty, use OCR
            try:
                if not pdf_text.strip():
                    print(f"No text found in {filename}. Using OCR...")
                    pdf_text = extract_text_from_images(file_path)
            except:
                print("not found")
            
            print()
            #print(pdf_text)
            
            year = extract_years(pdf_text)
            print()
            print()
            print()
            print()
            print("dates are")
            print(year)
            if(year== ['x']):
                str1="0-0"
            else:    
                #if(year)
                print(year[0])
                x=int(year[0])+1
                #print(x)
                str1=year[0]+"-"+str(x)
            
            print(str1)
            
            
            carrier=find_carrier_name(file_path)
            #print(carrier.strip())
            #lob=get_policy_type(file_path)
            str2= get_policy_type(file_path)
            if(str2):
                str2 = str2
            else:
                continue
            
            #str2=lob[0]
            if (str2=="Unknown"):
                str2="Unknown"
            print(str2)
            new_name=str1+" "+str2+"-Policy"+"-"+carrier+".pdf"
            print(new_name)
            new_file_path = os.path.join(folder_path, new_name)
            #if files of same name exists
            base, extension = os.path.splitext(new_file_path)
            counter = 1
            while os.path.exists(new_file_path):
                new_file_path = f"{base}_{counter}{extension}"
                counter += 1

            os.rename(file_path, new_file_path)

            #Record the old and new names in the Word document
            document.add_paragraph(f"Old Name: {filename} -> New Name: {new_name}")

    # Save the Word document
    doc_path = os.path.join(folder_path, "Old_vs_New_PDF_Names.docx")
    document.save(doc_path)
    print(f"Document saved to {doc_path}")

# def process_pdfs_in_folder(folder_path):
#     for filename in os.listdir(folder_path):
#         if(filename.endswith(".pdf")):
#             print(f"Found PDF: {filename}")
#             process_pdfs_in_folder(filename)

def main():
    # Specify the folder containing the PDF files
    # folder_path = "policies"
    # process_pdfs_in_folder(folder_path)

    exe_dir = os.path.dirname(os.path.abspath(sys.executable if hasattr(sys, 'frozen') else __file__))
    if hasattr(sys,'_MEIPASS'):
        exe_dir = sys._MEIPASS   
    else: 
        exe_dir = os.path.dirname(os.path.abspath(__file__))

    folder_path = os.path.join(exe_dir, "policies")
    print(folder_path)
    process_pdfs_in_folder(folder_path)


    
if __name__ == "__main__":
    main()
